import datetime
import os
import re
import shutil
import time

import qrcode
from docx import Document
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Cm, Pt
import  sys
import excel
import get_pic

img_file = r'dst'
filename = r'out.docx'

def generate_code(start, end):

    code_list = []
    for i in range(start, end+1):
        no_code = prefix + wuliaobianma + '_1500_'+ today0 + str(i).zfill(3)
        code_list.append(no_code)
    return code_list

def get_row_num(start,end):
    src = end - start +1
    return src





def create_table(tuopanhao,start, end,singe_list):
    document = Document()
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(9)
    data_count = get_row_num(start,end)
    if data_count % 2 != 0:
        data_count = data_count + 1
    row_num = 5 + int(data_count / 2)
    col_num = 4
    table = document.add_table(rows=row_num, cols=col_num, style='Table Grid')
    table = document.tables[-1]
    # 行宽
    for row in table.rows:
        row.cells[0].width = Cm(2)
    for row in table.rows:
        row.cells[1].width = Cm(6)
    for row in table.rows:
        row.cells[2].width = Cm(2)
    for row in table.rows:
        row.cells[3].width = Cm(6)

    # 第一行 物料编码
    table.cell(0, 0).text = '物料编码'
    table.cell(0,1).text = wuliaobianma
    # 第二行 数量
    table.cell(1, 0).text = '数量'
    table.cell(1, 1).text = shuliang
    # 第三行 托盘号
    table.cell(2, 0).text = '托盘号'
    table.cell(2, 1).text = tuopanhao
    table.rows[2].height = Cm(2.5)
    # 第四行 名称规格
    table.cell(3, 0).text = '名称规格'
    table.cell(3, 1).merge(table.cell(3, 3))
    table.cell(3, 1).text = mingchengguige
    # 第五行 No
    table.cell(4, 0).text = 'NO.'
    table.cell(4, 1).text = 'S.N.'
    table.cell(4, 2).text = 'NO.'
    table.cell(4, 3).text = 'S.N.'
    # 右上角二维码

    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_H,
        box_size=3,
        border=4
    )
    tuopanfile = img_file + '\\' + tuopanhao + '.jpg'
    qr.add_data(tuopanhao)
    qr.make(fit=True)
    img = qr.make_image()
    img.save(img_file + '\\' + tuopanhao + '.jpg')
    picture = table.cell(0, 2).paragraphs[0].add_run().add_picture(
        tuopanfile)
    picture.height = Cm(2.66)
    picture.width = Cm(2.66)

    table.cell(2, 2).paragraphs[0].add_run().add_text(tuopanhao)
    table.cell(0, 2).merge(table.cell(2, 3))

    # 序号与二维码
    content = singe_list
    for i in range(0, row_num - 5):
        table.cell(5 + i, 0).text = str(i * 2 + 1)

        picture = table.cell(5 + i, 1).paragraphs[0].add_run().add_picture(
            r'dst/' + content[i * 2] + '.jpg')
        picture.height = Cm(1.8)
        picture.width = Cm(1.8)

        table.cell(5 + i, 1).paragraphs[0].add_run().add_text('\n'+content[i * 2])


        if end != i * 2 + 1:
            table.cell(5 + i, 2).text = str(i * 2 + 2)
            picture = table.cell(5 + i, 3).paragraphs[0].add_run().add_picture(
                r'dst/' + content[i * 2 + 1] + '.jpg')
            picture.height = Cm(1.8)
            picture.width = Cm(1.8)
            table.cell(5 + i, 3).paragraphs[0].add_run().add_text('\n' + content[i * 2 + 1])
        else:
            print("奇数个")

    # 所有内容居中
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    document.save('A4绑拖单' + today0 +'-' +str(start)+'-'+str(end) + '.docx')

def main_program(start, end):
    if (os.path.exists('dst')):
        print('目录存在')
    else:
        os.mkdir('dst')
    list1 = generate_code(start, end)
    get_pic.get(list1)
    create_table(tuopanhao, start, end, list1)
    time.sleep(1)
    shutil.rmtree(r'dst')

if __name__ == "__main__":
    today = datetime.date.today()
    today0 = today.strftime('%y%m%d')
    today1 = today.strftime('%Y%m%d')
    prefix = '1011270'
    shuliang = '72000'
    wuliaobianma = '4JB9ES53022'
    mingchengguige = '模块二极管TPA4050S-2/150mil trench/含铅锡块0.30mm'
    print('python版本' + sys.version + '\n')
    step_1 = input("""请选择模式，输入数字后回车：
       1.固定模式：默认每48个编号出一个word，共出1-48，48-96，97-144三个word
       2.自定义模式：托盘号、编号可自定义
       """)
    if step_1 == '1':
        # tuopanhao = input("请输入托盘号:\n")
        for i in range(1,4):
            tuopanhao = 'JX'+today1+str(i).zfill(3)
            main_program((i-1)*48+1,i*48)
    elif step_1 == '2':
        tuopanhao = input('请输入托盘号，直接回车默认为'+'JX'+today1+'1'.zfill(3)+'\n')
        if len(tuopanhao) == 0:
            tuopanhao = 'JX'+today1+'1'.zfill(3)
        start = input('请输入开始编号，如1,直接回车默认为1\n')
        if len(start) == 0:
            start = 1
        start = int(start)
        end = input('请输入结束编号，如48,直接回车默认为48\n')
        if len(end) == 0:
            end = 48
        end = int(end)
        main_program(start,end)

